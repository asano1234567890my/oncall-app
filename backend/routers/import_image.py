"""画像・ファイルからスケジュール/医師名を読み取るAPI（Gemini）"""
from __future__ import annotations

import json
import os
import uuid
from datetime import date

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from pydantic import BaseModel
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession

from core.auth import get_current_hospital
from core.db import get_db
from models.doctor import Doctor
from models.shift import ShiftAssignment

router = APIRouter(prefix="/api/import", tags=["Import"])

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-3-flash-preview")

# ── Prompts ──────────────────────────────────────────────

SCHEDULE_PROMPT = """この画像は病院の当直表（オンコール表）です。
以下の情報をJSON形式で正確に抽出してください。

1. year: 年（4桁整数）
2. month: 月（1-12の整数）
3. shifts: 各日のシフト情報の配列

各シフトは以下の形式:
{
  "day": 日（1-31の整数）,
  "day_shift": 日直の医師名（ない場合はnull）,
  "night_shift": 当直の医師名（ない場合はnull）
}

重要なルール:
- 医師名は画像に書かれている通りに正確に抽出してください
- 日直・当直の区別がある場合はそれぞれ分けてください
- 日直・当直の区別がない場合は night_shift に入れてください
- 年月が画像から読み取れない場合は year: null, month: null としてください
- JSON以外のテキストは出力しないでください

出力形式（必ずこのJSON形式のみ）:
{"year": 2026, "month": 4, "shifts": [{"day": 1, "day_shift": null, "night_shift": "田中"}, ...]}"""

DOCTORS_PROMPT = """このファイルには医師（ドクター）の名前のリストが含まれています。
人名だけを正確に抽出してJSON配列で返してください。

重要なルール:
- 人名のみを抽出してください（役職・肩書き・診療科などは除く）
- 姓名がある場合はそのまま抽出してください（例: "田中太郎"）
- 姓だけの場合はそのまま抽出してください（例: "田中"）
- 重複は除いてください
- 人名以外の情報（日付、数値、場所名など）は含めないでください
- JSON以外のテキストは出力しないでください

出力形式（必ずこのJSON配列のみ）:
["田中", "佐藤花子", "鈴木"]"""


# ── Helpers ──────────────────────────────────────────────

def _extract_json(raw_text: str) -> str:
    """マークダウンコードブロックからJSON部分を抽出する。"""
    text = raw_text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        json_lines = []
        in_block = False
        for line in lines:
            if line.startswith("```") and not in_block:
                in_block = True
                continue
            if line.startswith("```") and in_block:
                break
            if in_block:
                json_lines.append(line)
        text = "\n".join(json_lines)
    return text


def _extract_text_from_file(content: bytes, mime_type: str, filename: str) -> str | None:
    """Excel/Word/PDFからテキストを抽出する。画像の場合はNone。"""
    lower = filename.lower()

    # Excel
    if lower.endswith((".xlsx", ".xls")) or "spreadsheet" in mime_type:
        import io
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                vals = [str(c) for c in row if c is not None]
                if vals:
                    lines.append("\t".join(vals))
        wb.close()
        return "\n".join(lines) if lines else None

    # Word
    if lower.endswith(".docx") or "wordprocessingml" in mime_type:
        import io
        from docx import Document
        doc = Document(io.BytesIO(content))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        # テーブルも抽出
        for table in doc.tables:
            for row in table.rows:
                vals = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if vals:
                    lines.append("\t".join(vals))
        return "\n".join(lines) if lines else None

    # PDF
    if lower.endswith(".pdf") or mime_type == "application/pdf":
        import io
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(content))
        lines = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                lines.append(text)
        return "\n".join(lines) if lines else None

    # テキスト
    if lower.endswith((".txt", ".csv")) or mime_type.startswith("text/"):
        return content.decode("utf-8", errors="replace")

    return None


def _call_gemini_vision(prompt: str, content: bytes, mime_type: str) -> str:
    """Gemini Vision APIを呼び出し、レスポンステキストを返す。"""
    from google import genai
    from google.genai import types

    client = genai.Client(api_key=GEMINI_API_KEY)
    response = client.models.generate_content(
        model=GEMINI_MODEL,
        contents=[
            types.Part.from_text(text=prompt),
            types.Part.from_bytes(data=content, mime_type=mime_type),
        ],
        config=types.GenerateContentConfig(temperature=0.1),
    )
    return response.text.strip()


def _call_gemini_text(prompt: str, text: str) -> str:
    """Gemini テキストAPIを呼び出し、レスポンステキストを返す。"""
    from google import genai
    from google.genai import types

    client = genai.Client(api_key=GEMINI_API_KEY)
    response = client.models.generate_content(
        model=GEMINI_MODEL,
        contents=[
            types.Part.from_text(text=f"{prompt}\n\n--- 以下がファイルの内容 ---\n{text}"),
        ],
        config=types.GenerateContentConfig(temperature=0.1),
    )
    return response.text.strip()


# ── 当直表 画像取込 ──────────────────────────────────────

ACCEPTED_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp",
                       ".xlsx", ".xls", ".docx", ".pdf", ".txt", ".csv"}

SCHEDULE_ACCEPTED_EXTENSIONS = ACCEPTED_EXTENSIONS


@router.post("/parse-image")
async def parse_image(
    file: UploadFile = File(...),
    hospital_id: uuid.UUID = Depends(get_current_hospital),
):
    """画像・Excel・Word・PDF・テキストからスケジュールデータを抽出する。"""
    if not GEMINI_API_KEY:
        raise HTTPException(status_code=500, detail="AI解析機能が設定されていません（APIキー未設定）")

    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="ファイルサイズが大きすぎます（上限10MB）")

    filename = file.filename or "unknown"
    mime_type = file.content_type or "application/octet-stream"

    ext = os.path.splitext(filename)[1].lower()
    if ext not in SCHEDULE_ACCEPTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail="対応していないファイル形式です。対応形式: 画像, Excel, Word, PDF, テキスト",
        )

    try:
        # ドキュメント系はテキスト抽出 → テキストAPI
        extracted = _extract_text_from_file(content, mime_type, filename)
        if extracted:
            raw_text = _call_gemini_text(SCHEDULE_PROMPT, extracted)
        elif mime_type.startswith("image/"):
            raw_text = _call_gemini_vision(SCHEDULE_PROMPT, content, mime_type)
        else:
            raise HTTPException(status_code=400, detail="ファイルの内容を読み取れませんでした")

        parsed = json.loads(_extract_json(raw_text))
    except json.JSONDecodeError:
        raise HTTPException(status_code=422, detail="ファイルの解析結果を正しく読み取れませんでした。別のファイルを試してください。")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"ファイル解析に失敗しました: {str(e)}")

    return parsed


class DoctorMapping(BaseModel):
    image_name: str
    doctor_id: str | None = None
    new_name: str | None = None


class ImportConfirmRequest(BaseModel):
    year: int
    month: int
    shifts: list[dict]
    doctor_mappings: list[DoctorMapping]


@router.post("/confirm")
async def confirm_import(
    body: ImportConfirmRequest,
    db: AsyncSession = Depends(get_db),
    hospital_id: uuid.UUID = Depends(get_current_hospital),
):
    """解析結果を確認し、スケジュールとしてDBに保存する。"""
    name_to_id: dict[str, uuid.UUID] = {}

    for mapping in body.doctor_mappings:
        if mapping.doctor_id:
            name_to_id[mapping.image_name] = uuid.UUID(mapping.doctor_id)
        elif mapping.new_name:
            import secrets
            new_doc = Doctor(
                hospital_id=hospital_id,
                name=mapping.new_name,
                access_token=secrets.token_urlsafe(32),
            )
            db.add(new_doc)
            await db.flush()
            name_to_id[mapping.image_name] = new_doc.id

    days_in_month = _days_in_month(body.year, body.month)
    start_date = date(body.year, body.month, 1)
    end_date = date(body.year, body.month, days_in_month)

    existing_doctors = await db.execute(
        select(Doctor.id).where(Doctor.hospital_id == hospital_id)
    )
    doc_ids = [row[0] for row in existing_doctors.all()]
    if doc_ids:
        await db.execute(
            delete(ShiftAssignment).where(
                ShiftAssignment.doctor_id.in_(doc_ids),
                ShiftAssignment.date >= start_date,
                ShiftAssignment.date <= end_date,
            )
        )

    saved_count = 0
    for shift in body.shifts:
        day = shift.get("day")
        if not day or day < 1 or day > days_in_month:
            continue

        shift_date = date(body.year, body.month, day)
        day_name = shift.get("day_shift")
        night_name = shift.get("night_shift")

        if day_name and day_name in name_to_id:
            db.add(ShiftAssignment(
                doctor_id=name_to_id[day_name],
                date=shift_date,
                shift_type="day",
            ))
            saved_count += 1

        if night_name and night_name in name_to_id:
            db.add(ShiftAssignment(
                doctor_id=name_to_id[night_name],
                date=shift_date,
                shift_type="night",
            ))
            saved_count += 1

    await db.commit()
    return {"message": f"スケジュールを保存しました（{saved_count}件）", "saved_count": saved_count}


# ── 医師名 一括取込 ──────────────────────────────────────


@router.post("/parse-doctors")
async def parse_doctors(
    file: UploadFile = File(...),
    hospital_id: uuid.UUID = Depends(get_current_hospital),
):
    """画像・Excel・Word・PDF・テキストから医師名リストを抽出する。"""
    if not GEMINI_API_KEY:
        raise HTTPException(status_code=500, detail="AI機能が設定されていません（APIキー未設定）")

    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="ファイルサイズが大きすぎます（上限10MB）")

    filename = file.filename or "unknown"
    mime_type = file.content_type or "application/octet-stream"

    # 拡張子チェック
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ACCEPTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"対応していないファイル形式です。対応形式: 画像, Excel, Word, PDF, テキスト",
        )

    try:
        # ドキュメント系はテキスト抽出 → テキストAPI
        extracted = _extract_text_from_file(content, mime_type, filename)
        if extracted:
            raw_text = _call_gemini_text(DOCTORS_PROMPT, extracted)
        else:
            # 画像 → Vision API
            if not mime_type.startswith("image/"):
                raise HTTPException(status_code=400, detail="ファイルの内容を読み取れませんでした")
            raw_text = _call_gemini_vision(DOCTORS_PROMPT, content, mime_type)

        names = json.loads(_extract_json(raw_text))
        if not isinstance(names, list):
            raise ValueError("not a list")
        # 文字列のみ、重複除去、空文字除去
        names = list(dict.fromkeys(str(n).strip() for n in names if str(n).strip()))
    except json.JSONDecodeError:
        raise HTTPException(status_code=422, detail="ファイルから医師名を読み取れませんでした。別のファイルを試してください。")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"ファイル解析に失敗しました: {str(e)}")

    return {"names": names}


class DoctorBulkRegisterRequest(BaseModel):
    names: list[str]


@router.post("/register-doctors")
async def register_doctors(
    body: DoctorBulkRegisterRequest,
    db: AsyncSession = Depends(get_db),
    hospital_id: uuid.UUID = Depends(get_current_hospital),
):
    """抽出した医師名を一括登録する。既存の同名医師はスキップ。"""
    import secrets

    # 既存医師名を取得
    existing = await db.execute(
        select(Doctor.name).where(Doctor.hospital_id == hospital_id)
    )
    existing_names = {row[0] for row in existing.all()}

    created = []
    skipped = []
    for name in body.names:
        name = name.strip()
        if not name:
            continue
        if name in existing_names:
            skipped.append(name)
            continue
        doc = Doctor(
            hospital_id=hospital_id,
            name=name,
            access_token=secrets.token_urlsafe(32),
        )
        db.add(doc)
        created.append(name)
        existing_names.add(name)

    await db.commit()
    return {
        "message": f"{len(created)}名を登録しました" + (f"（{len(skipped)}名は既に登録済み）" if skipped else ""),
        "created": created,
        "skipped": skipped,
    }


def _days_in_month(year: int, month: int) -> int:
    import calendar
    return calendar.monthrange(year, month)[1]
